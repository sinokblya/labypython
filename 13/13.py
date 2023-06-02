import docx
from guizero import App, Combo, Text, TextBox, PushButton
from materials import Oboi, Plitka, Laminate


def calculate_material():
    selected_material = combo_material.value
    area = float(txt_area.value)
    price = float(txt_price.value or 0)
    material_class = material_classes[selected_material]
    material_instance = material_class(price_per_unit=price)
    material_instance.calculate(area)
    txt_quantity.value = str(round(material_instance.quantity, 2))
    txt_total_price.value = str(round(material_instance.total_price, 2))


def save_report():
    selected_material = combo_material.value
    area = float(txt_area.value)
    price = float(txt_price.value or 0)
    material_class = material_classes[selected_material]
    material_instance = material_class(price_per_unit=price)
    material_instance.calculate(area)
    doc = docx.Document()
    doc.add_heading(f'Отчет по закупке "{selected_material}"', 0)
    doc.add_paragraph(f"Площадь помещения: {area} кв.м")
    doc.add_paragraph(f"Количество материала: {round(material_instance.quantity, 2)} шт.")
    doc.add_paragraph(f"Стоимость: {round(material_instance.total_price, 2)} руб.")
    doc.save(f'report_{selected_material.lower()}.docx')


app = App(title="Калькулятор для расчета отделочных материалов")
combo_material = Combo(app, options=["Обои", "Плитка", "Ламинат"])
txt_area = TextBox(app, text="")
txt_price = TextBox(app, text="")
txt_quantity = Text(app, text="Количество материала:")
txt_total_price = Text(app, text="Стоимость (руб.):")
btn_calculate = PushButton(app, text="Рассчитать", command=calculate_material)
btn_save = PushButton(app, text="Сохранить отчет", command=save_report)

material_classes = {"Обои": Oboi, "Плитка": Plitka, "Ламинат": Laminate}

app.display()
